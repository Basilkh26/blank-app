
# rtdp_live_dashboard_streamlit_emirates.py
# RTDP initiatives live dashboard (Emirates brand styling).
# Source: OneDrive Word doc; updates reflect on refresh.

import os
from typing import Optional, List, Tuple
from datetime import datetime

import numpy as np
import pandas as pd
import altair as alt
import streamlit as st
from docx import Document  # pip install python-docx
from docx.opc.exceptions import PackageNotFoundError
from streamlit.errors import StreamlitSecretNotFoundError
import html as html_module

# -----------------------------
# PAGE & BASIC CONFIG
# -----------------------------
st.set_page_config(page_title="RTDP Enhancements Dashboard — Emirates", layout="wide")

DOC_PATH = r"C:\Users\S435603\OneDrive - Emirates Group\Desktop\RTDP priority\RTDP new initiatives.docx"
LOGO_DIR = r"C:\Users\S435603\OneDrive - Emirates Group\Desktop\RTDP priority"
RELEASE_NOTE_FILE = os.path.join(LOGO_DIR, "Release dates.docx")  # note file in same folder

# Optional secrets helper (no broad exceptions)
def _get_secret(name: str, default: Optional[str]) -> Optional[str]:
    """Safely read a secret key; return default if secrets.toml is absent or key missing."""
    try:
        val = st.secrets.get(name, default)  # may raise StreamlitSecretNotFoundError if secrets.toml missing
        return None if val is None else str(val)
    except StreamlitSecretNotFoundError:
        return default
    except (AttributeError, TypeError):
        return default

EXCEL_PATH = _get_secret("EXCEL_PATH", "RTDP_Dashboard.xlsx")
LOGO_PATH  = _get_secret("LOGO_PATH", None)

# Env overrides (optional)
EXCEL_PATH = os.environ.get("RTDP_EXCEL_PATH", EXCEL_PATH)
LOGO_PATH  = os.environ.get("RTDP_LOGO_PATH", LOGO_PATH)

# -----------------------------
# BRAND COLORS
# -----------------------------
PRIMARY_RED = "#D71A21"     # Emirates red
ACCENT_GOLD = "#FEF8D5"
TEXT_DARK   = "#333333"
GREEN_OK    = "#2ECC71"     # Green for 'Delivered'
AMBER_DEV   = "#F1C40F"     # Amber for 'In Development'
BLUE_PLAN   = "#3498DB"     # Blue for 'Planned'

# -----------------------------
# BRAND CSS
# -----------------------------
st.markdown(
    f"""
<style>
/* Headings */
h1, h2, h3, h4, h5, h6 {{
  color: {PRIMARY_RED};
  font-family: Georgia, 'Times New Roman', serif;
}}
/* KPI cards */
div[data-testid="stMetric"] > div {{
  background:{ACCENT_GOLD};
  padding:12px;
  border-radius:8px;
  border:1px solid rgba(215,26,33,0.15);
}}
/* Risk/info box */
.emirates-note {{
  border-left: 4px solid {PRIMARY_RED};
  background:#ffffff;
  padding:10px;
  border-radius:6px;
}}
/* Optional tighter top padding */
.block-container {{
  padding-top: 0.75rem;
}}
/* Fix table column widths + prevent wrapping in Status column */
table.emirates-table {{
  width: 100%;
  border-collapse: collapse;
  table-layout: fixed;
}}
table.emirates-table th, table.emirates-table td {{
  padding: 6px;
  text-align: left;
  vertical-align: top;
}}
table.emirates-table colgroup col.col-init     {{ width: 24%; }}
table.emirates-table colgroup col.col-status   {{ width: 20%; }}
table.emirates-table colgroup col.col-release  {{ width: 16%; }}
table.emirates-table colgroup col.col-purpose  {{ width: 20%; }}
table.emirates-table colgroup col.col-solution {{ width: 20%; }}
table.emirates-table td.status-cell {{
  white-space: nowrap; /* keep "In Development" on one line */
}}
</style>
""",
    unsafe_allow_html=True,
)

# -----------------------------
# UTILITIES
# -----------------------------
def find_logo_in_dir(folder: str) -> Optional[str]:
    """Find the first likely logo/image in the given directory."""
    if not os.path.isdir(folder):
        return None
    exts = {".png", ".jpg",".svg", ".jpeg", ".gif", ".bmp", ".webp"}
    names = sorted(os.listdir(folder))
    # prefer likely names first
    for prefix in ("emirates", "logo", "brand", "emirates_logo"):
        for f in names:
            if f.lower().startswith(prefix) and os.path.splitext(f)[1].lower() in exts:
                return os.path.join(folder, f)
    # fallback: any image
    for f in names:
        if os.path.splitext(f)[1].lower() in exts:
            return os.path.join(folder, f)
    return None

def read_from_word(path: str) -> pd.DataFrame:
    """Read the first table from the Word doc into a DataFrame."""
    if not os.path.exists(path):
        raise FileNotFoundError(f"Word file not found: {path}")
    doc = Document(path)
    if not doc.tables:
        raise ValueError("No tables found in Word document.")
    table = doc.tables[0]
    rows: List[List[str]] = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        if i == 0:  # header
            continue
        if not any(cells):
            continue
        if cells[0].lower().startswith("initiative"):
            continue
        rows.append(cells)
    cols = ['Initiative', 'Purpose / Impact', 'Solution', 'Status', 'Target Release']
    return pd.DataFrame(rows, columns=cols)

def parse_release_mmyy(token: str) -> Tuple[int, int]:
    """Parse a 'Dec25' style string to (year, month). Unknown → (9999, 99)."""
    month_order = {"Jan":1,"Feb":2,"Mar":3,"Apr":4,"May":5,"Jun":6,"Jul":7,"Aug":8,"Sep":9,"Oct":10,"Nov":11,"Dec":12}
    if isinstance(token, str) and len(token) >= 5 and token[:3] in month_order:
        try:
            y = 2000 + int(token[3:5])
            m = month_order[token[:3]]
            return y, m
        except (ValueError, TypeError):
            return 9999, 99
    return 9999, 99

def parse_version(token: str) -> Tuple[int, int, int]:
    """Parse 'v8.1.0' to (8,1,0) for sorting. Non-version → (9999,99,99)."""
    try:
        if token.startswith('v'):
            parts = token[1:].split('.')
            nums = [int(p) for p in parts]
            while len(nums) < 3:
                nums.append(0)
            return nums[0], nums[1], nums[2]
    except (ValueError, AttributeError):
        pass
    return 9999, 99, 99

def release_label(tr: str) -> str:
    """
    Label to show on chart:
      - For versions: keep 'v8.x.x'
      - For month-year tokens: keep as 'Dec25'
      - For Delivered bucket: 'Delivered (v8.x)'
    """
    if tr is None:
        return "Unknown"
    t = str(tr).strip()
    if t.startswith('v'):
        return t
    return t  # Dec25, Mar26, etc.

def build_release_domain(series: pd.Series) -> List[str]:
    """
    Custom x-axis order (by "release number"):
      1) 'Delivered (v8.x)' first (if present),
      2) Versions 'vX.Y.Z' ascending,
      3) Month-year tokens ascending (chronological).
    """
    items = list(pd.unique(series.dropna()))
    ordered_domain: List[str] = []

    # 1) Delivered first
    if "Delivered (v8.x)" in items:
        ordered_domain.append("Delivered (v8.x)")
        items.remove("Delivered (v8.x)")

    # 2) Versions ascending
    versions = [i for i in items if isinstance(i, str) and i.startswith('v')]
    others   = [i for i in items if i not in versions]
    versions_sorted = sorted(versions, key=parse_version)
    ordered_domain.extend(versions_sorted)

    # 3) Month-year tokens chronological
    months_sorted = sorted(others, key=parse_release_mmyy)
    ordered_domain.extend(months_sorted)

    return ordered_domain

def prepare(df: pd.DataFrame) -> pd.DataFrame:
    """Standardize fields and add ReleaseLabel/ReleaseGroup/ReleaseDate and RiskFlag."""
    out = df.copy()
    out['Status'] = out['Status'].str.strip()
    out['RiskFlag'] = out['Target Release'].str.contains('Escalated', case=False, na=False)
    out['ReleaseLabel'] = out['Target Release'].map(release_label)

    # Safe date mapping only when parseable
    def to_date(x):
        if pd.isna(x) or str(x).startswith('v'):
            return None
        y, m = parse_release_mmyy(x)
        return datetime(y, m, 1) if y != 9999 and m != 99 else None

    out['ReleaseGroup'] = out['ReleaseLabel']
    out['ReleaseDate'] = out['Target Release'].map(to_date)
    return out

def read_release_notes(note_path: str) -> Optional[str]:
    """Read paragraphs and first table from 'Release dates.docx' into HTML note."""
    if not os.path.exists(note_path):
        return None
    try:
        doc = Document(note_path)
        parts: List[str] = []
        # paragraphs
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                parts.append(html_module.escape(txt))
        # first table (if any)
        if doc.tables:
            tbl = doc.tables[0]
            tbl_rows_html: List[str] = []
            for r in tbl.rows:
                cells = [c.text.strip() for c in r.cells]
                tr_html = "<tr>" + "".join(
                    f"<td style='padding:4px;border:1px solid #ddd;'>{html_module.escape(cell)}</td>"
                    for cell in cells
                ) + "</tr>"
                tbl_rows_html.append(tr_html)
            parts.append(
                "<table style='border-collapse:collapse;margin-top:6px;'>"
                + "".join(tbl_rows_html)
                + "</table>"
            )
        return "<br/>".join(parts)
    except (PackageNotFoundError, OSError, ValueError):
        # If the Word file is missing/corrupt/unreadable, show a helpful message.
        return "Release dates are maintained in 'Release dates.docx' in the RTDP priority folder."

# -----------------------------
# HEADER (spacer + logo + title) — smaller logo, no yellow deprecation note
# -----------------------------
# Pull down the header by ~1 cm
st.markdown("<div style='height:1cm'></div>", unsafe_allow_html=True)

resolved_logo = LOGO_PATH if LOGO_PATH and os.path.exists(LOGO_PATH) else find_logo_in_dir(LOGO_DIR)

c1, c2 = st.columns([0.7, 6.3])  # narrow logo column
with c1:
    if resolved_logo:
        # Fixed pixel width (no use_column_width) → no deprecation note
        st.image(resolved_logo, width=120)
    else:
        st.write("")  # spacer

with c2:
    st.title("RTDP Enhancements Dashboard")
    st.caption(f"Updated: {datetime.now().strftime('%d %b %Y %H:%M')} • Source: Word document (OneDrive)")

# Refresh button (supported API)
if st.sidebar.button("Refresh now"):
    st.rerun()

# -----------------------------
# DATA LOAD
# -----------------------------
try:
    source_df = read_from_word(DOC_PATH)
except (FileNotFoundError, ValueError, OSError) as e:
    st.warning(f"{e}\nFalling back to Excel if available…")
    try:
        source_df = pd.read_excel(EXCEL_PATH, sheet_name='Initiatives', engine='openpyxl')
    except (FileNotFoundError, ValueError, OSError, ImportError) as e2:
        st.error(f"Failed to load data: {e2}")
        st.stop()

data_df = prepare(source_df)

# -----------------------------
# KPIs
# -----------------------------
kc1, kc2, kc3, kc4 = st.columns(4)
with kc1:
    st.metric("Total Initiatives", len(data_df))
with kc2:
    st.metric("Delivered", int(np.count_nonzero(data_df['Status'].eq('Delivered'))))
with kc3:
    st.metric("In Development", int(np.count_nonzero(data_df['Status'].eq('In Development'))))
with kc4:
    st.metric("Planned", int(np.count_nonzero(data_df['Status'].eq('Planned'))))

st.divider()

# -----------------------------
# FILTERS
# -----------------------------
status_sel  = st.multiselect(
    "Filter by Status",
    sorted(data_df['Status'].dropna().unique()),
    default=list(sorted(data_df['Status'].dropna().unique())),
)
release_sel = st.multiselect(
    "Filter by Release Number",
    sorted(pd.Series(data_df['ReleaseLabel']).dropna().unique()),
    default=list(sorted(pd.Series(data_df['ReleaseLabel']).dropna().unique())),
)
search_text = st.text_input("Search initiative (contains)", "")

filtered_df = data_df[data_df['Status'].isin(status_sel) & data_df['ReleaseLabel'].isin(release_sel)]
if search_text:
    filtered_df = filtered_df[filtered_df['Initiative'].str.contains(search_text, case=False, na=False)]

# -----------------------------
# RELEASE MIX CHART (by Release Number, Delivered first)
# -----------------------------
release_domain = build_release_domain(filtered_df['ReleaseLabel'])
rel_counts = filtered_df.groupby('ReleaseLabel').size().reset_index(name='Count')

chart = alt.Chart(rel_counts).mark_bar(color=PRIMARY_RED).encode(
    x=alt.X('ReleaseLabel:N',
            title='Release Number',
            scale=alt.Scale(domain=release_domain)),  # Delivered first, then versions, then months
    y=alt.Y('Count:Q', title='Count'),
    tooltip=['ReleaseLabel', 'Count'],
).properties(width=720, height=340, title='Initiatives by Release Number')
st.altair_chart(chart, use_container_width=True)

# -----------------------------
# DETAIL TABLE (Delivered = GREEN badge; Status column wider + nowrap)
# -----------------------------
st.subheader("Initiatives Status Detail")

badge_map = {
    'Delivered':      f"<span style='background:{GREEN_OK};color:white;padding:4px 8px;border-radius:12px;'>Delivered</span>",
    'In Development': f"<span style='background:{AMBER_DEV};color:black;padding:4px 8px;border-radius:12px;'>In Development</span>",
    'Planned':        f"<span style='background:{BLUE_PLAN};color:white;padding:4px 8px;border-radius:12px;'>Planned</span>",
}
display_df = filtered_df.copy()
display_df['Status Badge'] = display_df['Status'].map(lambda s: badge_map.get(s, s))

def row_html(row):
    # tint delivered rows very lightly
    bg = "background-color: rgba(46, 204, 113, 0.10);" if row['Status'] == 'Delivered' else ""
    return (
        f"<tr style='{bg}'>"
        f"<td>{row['Initiative']}</td>"
        f"<td class='status-cell'>{row['Status Badge']}</td>"
        f"<td>{row['ReleaseLabel']}</td>"
        f"<td>{row['Purpose / Impact']}</td>"
        f"<td>{row['Solution']}</td>"
        f"</tr>"
    )

table_rows_html = "\n".join([row_html(r) for _, r in display_df.iterrows()])
table_html = (
    "<table class='emirates-table'>"
    "<colgroup>"
    "<col class='col-init' />"
    "<col class='col-status' />"
    "<col class='col-release' />"
    "<col class='col-purpose' />"
    "<col class='col-solution' />"
    "</colgroup>"
    f"<thead><tr style='background-color:{PRIMARY_RED}; color:white;'>"
    "<th>Initiative</th><th>Status</th><th>Release</th><th>Purpose / Impact</th><th>Solution</th>"
    "</tr></thead>"
    f"<tbody>{table_rows_html}</tbody>"
    "</table>"
)
st.markdown(table_html, unsafe_allow_html=True)

# -----------------------------
# RELEASE DATES NOTE (from 'Release dates.docx')
# -----------------------------
note_html = read_release_notes(RELEASE_NOTE_FILE)
if note_html:
    st.markdown("<div class='emirates-note'><strong>Release Dates:</strong><br/>" + note_html + "</div>", unsafe_allow_html=True)
else:
    st.markdown("<div class='emirates-note'>Release dates file not found in the RTDP priority folder.</div>", unsafe_allow_html=True)

# -----------------------------
# EXPORT
# -----------------------------
st.download_button(
    label="Download filtered data (CSV)",
    data=filtered_df.to_csv(index=False).encode('utf-8'),
    file_name='RTDP_filtered.csv',
    mime='text/csv',
)
